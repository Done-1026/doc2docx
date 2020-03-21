import subprocess
import shutil
import os
import logging
import re

import docx

logging.basicConfig(level=logging.INFO)


class OptDoc:
    """
    因无法直接对doc文件进行修改，所以
    1、doc 转换成 Docx
    2、修改docx文件内容

    """
    default_path = os.getcwd()

    def __init__(self, docx_path=default_path+'_docx', dst_path=default_path+'_new',
                 doc_path=default_path, copy_all=False):
        self.doc_path = os.path.abspath(doc_path)
        self.docx_path = os.path.abspath(docx_path)
        self.dst_path = os.path.abspath(dst_path)
        self.copy_all = copy_all
        self.ignore = ['LibreOffice.app']
        if not os.path.exists(self.docx_path):
            os.mkdir(self.docx_path)
        if not os.path.exists(self.dst_path):
            os.mkdir(self.dst_path)

    @staticmethod
    def _doc_to_docx_on_unix(filename, path):
        """
        mac os系统下需要下载libreoffice,再使用命令行转换文件
        :param filename: 要转换的doc文件的名称，带上后缀
        :return:
        """
        office_path = r'/Applications/LibreOffice.app/Contents/MacOS/soffice'
        cmd = office_path + ' --headless --convert-to docx ' + filename + ' --outdir ' + path
        logging.info(cmd)
        p = subprocess.call(cmd, shell=True)

    def _doc_to_docx_on_win(self):
        pass

    def doc_to_docx(self):
        """
        将doc_path目录下的所有doc文件转换为docx,并以相同目录树存入_docx的文件夹中，非doc格式文件，进行复制操作
        """
        for root, _, files in os.walk(self.doc_path):
            path = root.replace(self.doc_path, self.docx_path)
            os.chdir(root)
            if not os.path.exists(path):
                os.mkdir(path)
            #logging.info(files)
            filename = ' '.join(re.findall(r'\S+\.docx?\s', ' '.join(files)+' '))
            #logging.info(filename)
            if self.copy_all:
                # 如果copy_all=True,将非doc/docx的文件也复制到新的目录中
                for i in files:
                    if not i.endswith('.doc') and not i.endswith('.docx'):
                        try:
                            shutil.copy(i, path)
                        except PermissionError as f:
                            # 需要替换的文件无w权限时，修改文件的权限
                            os.chmod(f.filename, int('0o100644', 8))
                            shutil.copy(i, path)
            if filename.strip():
                # 如果该目录中存在doc/docx后缀的文件，则执行以下操作，否则进入下个循环
                if os.name is 'posix':
                    self._doc_to_docx_on_unix(filename, path)
                else:
                    pass

    @staticmethod
    def handle_document(src_docx, dst_docx):
        """
        :param src_docx: 源docx文件的docx.Document对象
        :param dst_docx: 内容进行修改后的docx.Document对像
        """
        s = ''
        for i in src_docx.paragraphs:
            if not i.text.isspace():
                s += i.text + '\n'
        #logging.info(s)
        q, a = s.split("答案部分")
        part_re = re.compile(r'''[一二三四五六七八九十]{1,2}、[A-Z]\d?.*''')
        q_a_re = re.compile(r'''(?<=\n)\d{1,3}、.*?(?=\d{1,3}、)''', re.DOTALL)
        first_part = part_re.split(q)
        title, q_part = first_part[0], first_part[1:]
        title = dst_docx.add_heading(title, 1)
        parts_title = part_re.findall(q)
        a_part = part_re.split(a)[1:]

        for i in range(len(q_part)):
            dst_docx.add_paragraph(parts_title[i])
            questions = q_a_re.findall(q_part[i]+'999、')
            answers = q_a_re.findall(a_part[i]+'999、')
            for j in range(len(questions)):
                try:
                    dst_docx.add_paragraph(questions[j]+answers[j])
                except:
                    global x
                    x = questions
                    global y
                    y = answers
                    raise

    def modify_docx(self):
        """
        修改_docx文件夹下的docx文件内容，并以相同目录树结构存入_new文件夹中，非docx格式的进行复制操作
        :return:
        """
        for root, _, files in os.walk(self.docx_path):
            new_path = root.replace(self.docx_path, self.dst_path)
            if not os.path.exists(new_path):
                os.mkdir(new_path)
            if files:
                for filename in files:
                    if filename.endswith('.docx'):
                        docx_obj = docx.Document(os.path.join(root, filename))
                        new_obj = docx.Document()
                        try:
                            self.handle_document(docx_obj, new_obj)
                        except ValueError:
                            print('文件名为:%s的文件内容不符合格式' % filename)
                            continue
                        new_obj.styles['Normal'].font.name = u'黑体'
                        dst_abspath = os.path.join(new_path, filename)
                        logging.info(dst_abspath)
                        new_obj.save(dst_abspath)
                    else:
                        if self.copy_all:
                            no_docx_path = os.path.join(root, filename)
                            try:
                                shutil.copy(no_docx_path, new_path)
                            except PermissionError as f:
                                os.chmod(f.filename, int('0o100644', 8))
                                shutil.copy(no_docx_path, new_path)
                        else:
                            pass


if __name__ == '__main__':
    print('注：默认情况下，该程序将影响它所在的当前目录，以下我们将该目录称为s目录，而它的上级目录称为p目录,如需指定目录，请参考第2条\n'
          '1、将s目录下的所有doc文件转换为docx，并存放在p目录中名为<s目录名_docx>的目录中，我们称为b1目录;'
          '并将b1目录中内容进行处理，结果存放在p目录中名为<b1目录名_new>的目录中;\n'
          '2、若想自由指定源目录及其他操作，请选择此项!')

    while True:
        client = input('>请依照以上说明，输入相应编号:')
        if client == '1':
            o = OptDoc()
            break
        elif client == '2':
            while True:
                docx_path_s = input('>请输入包含需要修改的docx文档内容的目录路径：')
                if os.path.isdir(docx_path_s):
                    break
                else:
                    print('>您输入的路径不正确，请重新输入！')
                    continue
            while True:
                dst_path_s = input('>请输入存放修改后的docx文档的目录路径：')
                if os.path.isdir(dst_path_s):
                    break
                else:
                    print('>您输入的路径不正确，请重新输入！')
                    continue
            while True:
                doc_path_s = input('>请输入需要进行doc转换为docx的源文件目录路径：(若不进行doc到docx的转换可回车跳过)')
                if os.path.isdir(doc_path_s) or doc_path_s == '':
                    break
                else:
                    print('>您输入的路径不正确，请重新输入！')
                    continue
            while True:
                copy_all_input = input('>是否复制非doc、docx格式的文件:(Y/N)')
                if copy_all_input not in ['Y', 'y', 'N', 'n']:
                    print('>输入错误，请重新输入！')
                    continue
                else:
                    copy_all_file = True if copy_all_input in ['Y', 'y'] else False
                    break
            o = OptDoc(docx_path=docx_path_s, dst_path=dst_path_s, doc_path=doc_path_s, copy_all=copy_all_file)
            break
        else:
            print('>输入有误，请重新输入!')
    while True:
        whether_doc_to_docx = input('>是否进行doc到docx的转换，如已经为docx文件，则不需转换!(Y/N)')
        if whether_doc_to_docx in ['Y', 'y']:
            o.doc_to_docx()
            break
        elif whether_doc_to_docx in ['N', 'n']:
            break
        else:
            print('>输入错误，请重新输入！')
    o.modify_docx()



